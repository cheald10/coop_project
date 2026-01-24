# app/services/coop_plan.py
"""
Service module for generating COOP plans (Word + PDF)
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
from django.conf import settings
from django.core.files import File
from app.models import (
    Division, EssentialFunction, CriticalApplication, KeyPersonnel,
    VitalRecord, Dependency, AlternateFacility, Communication,
    RecoveryPriority, DivisionMetadata, GeneratedPlan
)


def generate_coop_plan_for_division(division_id):
    """
    Main function to generate a complete COOP plan document.
    
    Returns:
        dict: {
            'success': bool,
            'plan': GeneratedPlan instance,
            'docx_path': str,
            'pdf_path': str,
            'errors': list
        }
    """
    try:
        division = Division.objects.get(pk=division_id)
    except Division.DoesNotExist:
        return {
            'success': False,
            'errors': [f'Division with ID {division_id} not found']
        }
    
    # Create Word document
    doc = Document()
    
    # Add cover page
    _add_cover_page(doc, division)
    
    # Add table of contents placeholder
    doc.add_page_break()
    doc.add_heading('Table of Contents', 0)
    doc.add_paragraph('[Table of Contents will be generated automatically]')
    
    # Add division metadata
    doc.add_page_break()
    _add_division_metadata_section(doc, division)
    
    # Add essential functions
    doc.add_page_break()
    _add_essential_functions_section(doc, division)
    
    # Add critical applications
    doc.add_page_break()
    _add_critical_applications_section(doc, division)
    
    # Add key personnel
    doc.add_page_break()
    _add_key_personnel_section(doc, division)
    
    # Add vital records
    doc.add_page_break()
    _add_vital_records_section(doc, division)
    
    # Add dependencies
    doc.add_page_break()
    _add_dependencies_section(doc, division)
    
    # Add alternate facilities
    doc.add_page_break()
    _add_alternate_facilities_section(doc, division)
    
    # Add communications
    doc.add_page_break()
    _add_communications_section(doc, division)
    
    # Add recovery priorities
    doc.add_page_break()
    _add_recovery_priorities_section(doc, division)
    
    # Save Word document
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename_base = f"COOP_Plan_{division.name.replace(' ', '_')}_{timestamp}"
    
    docx_filename = f"{filename_base}.docx"
    docx_path = os.path.join(settings.MEDIA_ROOT, 'coop_plans', 'docx', docx_filename)
    
    # Ensure directory exists
    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    
    doc.save(docx_path)
    
    # Convert to PDF (requires LibreOffice or docx2pdf)
    pdf_path = _convert_to_pdf(docx_path)
    
    # Create GeneratedPlan record
    new_version = division.plan_version + 1
    
    with open(docx_path, 'rb') as docx_file:
        with open(pdf_path, 'rb') as pdf_file:
            plan = GeneratedPlan.objects.create(
                division=division,
                version=new_version,
                docx_file=File(docx_file, name=os.path.basename(docx_path)),
                pdf_file=File(pdf_file, name=os.path.basename(pdf_path)),
                notes=f"Auto-generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            )
    
    # Update division version
    division.plan_version = new_version
    division.save()
    
    return {
        'success': True,
        'plan': plan,
        'docx_path': docx_path,
        'pdf_path': pdf_path,
        'errors': []
    }


def _add_cover_page(doc, division):
    """Add cover page to document"""
    title = doc.add_heading(level=0)
    title.text = 'CONTINUITY OF OPERATIONS PLAN'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    subtitle = doc.add_heading(division.name, level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    metadata_para = doc.add_paragraph()
    metadata_para.add_run(f"Version: {division.plan_version}\n").bold = True
    metadata_para.add_run(f"Status: {division.plan_status}\n")
    metadata_para.add_run(f"Last Updated: {division.last_updated.strftime('%B %d, %Y')}\n")
    metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_division_metadata_section(doc, division):
    """Add division profile/metadata section"""
    doc.add_heading('Division Profile', 1)
    
    try:
        metadata = division.divisionmetadata
        
        doc.add_heading('Director', 2)
        doc.add_paragraph(metadata.director or 'Not specified')
        
        doc.add_heading('Mission Statement', 2)
        doc.add_paragraph(metadata.mission_statement or 'Not specified')
        
        doc.add_heading('Primary Location', 2)
        doc.add_paragraph(metadata.primary_location or 'Not specified')
        
        doc.add_heading('Staff Count', 2)
        doc.add_paragraph(str(metadata.staff_count))
        
        doc.add_heading('Hours of Operation', 2)
        doc.add_paragraph(metadata.hours_of_operation or 'Not specified')
        
    except DivisionMetadata.DoesNotExist:
        doc.add_paragraph('Division metadata not yet configured.')


def _add_essential_functions_section(doc, division):
    """Add essential functions section"""
    doc.add_heading('Essential Functions', 1)
    
    functions = EssentialFunction.objects.filter(division=division).order_by('priority', 'name')
    
    if not functions.exists():
        doc.add_paragraph('No essential functions defined.')
        return
    
    for func in functions:
        doc.add_heading(func.name, 2)
        doc.add_paragraph(f"Priority: {func.priority}")
        doc.add_paragraph(f"Owner: {func.owner}")
        doc.add_paragraph(f"MTD: {func.mtd}")
        doc.add_paragraph(f"RTO: {func.rto}")
        
        doc.add_heading('Description', 3)
        doc.add_paragraph(func.description or 'Not specified')
        
        doc.add_heading('Dependencies', 3)
        doc.add_paragraph(func.dependencies or 'None identified')
        
        doc.add_heading('Alternate Procedures', 3)
        doc.add_paragraph(func.alternate_procedures or 'None defined')
        
        doc.add_paragraph()  # Spacing


def _add_critical_applications_section(doc, division):
    """Add critical applications section"""
    doc.add_heading('Critical Applications', 1)
    
    apps = CriticalApplication.objects.filter(division=division).order_by('recovery_tier', 'name')
    
    if not apps.exists():
        doc.add_paragraph('No critical applications defined.')
        return
    
    # Create table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Application'
    hdr_cells[1].text = 'Hosting'
    hdr_cells[2].text = 'Recovery Tier'
    hdr_cells[3].text = 'RTO'
    hdr_cells[4].text = 'Vendor Contact'
    
    # Data rows
    for app in apps:
        row_cells = table.add_row().cells
        row_cells[0].text = app.name
        row_cells[1].text = app.hosting_environment
        row_cells[2].text = app.recovery_tier
        row_cells[3].text = app.rto
        row_cells[4].text = app.vendor_contact or 'N/A'


def _add_key_personnel_section(doc, division):
    """Add key personnel section"""
    doc.add_heading('Key Personnel', 1)
    
    personnel = KeyPersonnel.objects.filter(division=division).order_by('primary_or_alternate', 'role')
    
    if not personnel.exists():
        doc.add_paragraph('No key personnel defined.')
        return
    
    # Create table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Role'
    hdr_cells[2].text = 'Type'
    hdr_cells[3].text = 'Phone'
    hdr_cells[4].text = 'Email'
    
    # Data rows
    for person in personnel:
        row_cells = table.add_row().cells
        row_cells[0].text = person.name
        row_cells[1].text = person.role
        row_cells[2].text = person.primary_or_alternate
        row_cells[3].text = person.mobile_phone or person.work_phone or 'N/A'
        row_cells[4].text = person.email or 'N/A'


def _add_vital_records_section(doc, division):
    """Add vital records section"""
    doc.add_heading('Vital Records', 1)
    
    records = VitalRecord.objects.filter(division=division).order_by('priority', 'name')
    
    if not records.exists():
        doc.add_paragraph('No vital records defined.')
        return
    
    for record in records:
        doc.add_heading(record.name, 2)
        doc.add_paragraph(f"Type: {record.record_type}")
        doc.add_paragraph(f"Priority: {record.priority}")
        doc.add_paragraph(f"Format: {record.format}")
        doc.add_paragraph(f"Storage Location: {record.storage_location}")
        doc.add_paragraph(f"Backup Location: {record.backup_location}")
        doc.add_paragraph()


def _add_dependencies_section(doc, division):
    """Add dependencies section"""
    doc.add_heading('Dependencies', 1)
    
    deps = Dependency.objects.filter(division=division).order_by('criticality', 'name')
    
    if not deps.exists():
        doc.add_paragraph('No dependencies defined.')
        return
    
    for dep in deps:
        doc.add_heading(dep.name, 2)
        doc.add_paragraph(f"Type: {dep.dependency_type}")
        doc.add_paragraph(f"Criticality: {dep.criticality}")
        doc.add_paragraph(f"Description: {dep.description or 'N/A'}")
        doc.add_paragraph(f"Vendor Contact: {dep.vendor_contact or 'N/A'}")
        doc.add_paragraph()


def _add_alternate_facilities_section(doc, division):
    """Add alternate facilities section"""
    doc.add_heading('Alternate Facilities', 1)
    
    facilities = AlternateFacility.objects.filter(division=division)
    
    if not facilities.exists():
        doc.add_paragraph('No alternate facilities defined.')
        return
    
    for facility in facilities:
        doc.add_heading(facility.name, 2)
        doc.add_paragraph(f"Type: {facility.facility_type}")
        doc.add_paragraph(f"Address: {facility.address}")
        doc.add_paragraph(f"Capacity: {facility.capacity} persons")
        doc.add_paragraph(f"IT Availability: {facility.it_availability}")
        doc.add_paragraph(f"Contact: {facility.contact or 'N/A'}")
        doc.add_paragraph()


def _add_communications_section(doc, division):
    """Add communications section"""
    doc.add_heading('Communications', 1)
    
    comms = Communication.objects.filter(division=division)
    
    if not comms.exists():
        doc.add_paragraph('No communications defined.')
        return
    
    for comm in comms:
        doc.add_heading(comm.communication_type, 2)
        doc.add_paragraph(f"Method: {comm.method}")
        doc.add_paragraph(f"Primary Contact: {comm.primary_contact}")
        doc.add_paragraph(f"Backup Contact: {comm.backup_contact or 'N/A'}")
        doc.add_paragraph(f"Details: {comm.contact_details}")
        doc.add_paragraph()


def _add_recovery_priorities_section(doc, division):
    """Add recovery priorities section"""
    doc.add_heading('Recovery Priorities', 1)
    
    priorities = RecoveryPriority.objects.filter(division=division).order_by('priority_level')
    
    if not priorities.exists():
        doc.add_paragraph('No recovery priorities defined.')
        return
    
    # Create table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Priority Level'
    hdr_cells[1].text = 'Item'
    hdr_cells[2].text = 'Type'
    hdr_cells[3].text = 'Rationale'
    
    # Data rows
    for priority in priorities:
        row_cells = table.add_row().cells
        row_cells[0].text = str(priority.priority_level)
        row_cells[1].text = priority.item_name
        row_cells[2].text = priority.item_type
        row_cells[3].text = priority.rationale or 'N/A'


def _convert_to_pdf(docx_path):
    """
    Convert Word document to PDF.
    
    Options:
    1. Use docx2pdf library (Windows/Mac with Word installed)
    2. Use LibreOffice in headless mode (Linux/cross-platform)
    
    Returns path to PDF file.
    """
    pdf_path = docx_path.replace('.docx', '.pdf').replace('/docx/', '/pdf/')
    
    # Ensure PDF directory exists
    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
    
    try:
        # Try using docx2pdf first
        from docx2pdf import convert
        convert(docx_path, pdf_path)
    except ImportError:
        # Fall back to LibreOffice
        import subprocess
        subprocess.run([
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', os.path.dirname(pdf_path),
            docx_path
        ], check=True)
        
        # LibreOffice creates PDF in same directory, move if needed
        temp_pdf = docx_path.replace('.docx', '.pdf')
        if temp_pdf != pdf_path and os.path.exists(temp_pdf):
            os.rename(temp_pdf, pdf_path)
    
    return pdf_path
