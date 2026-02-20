from dbos import workflow, step

@workflow()
def process_coop_request(user_id: int, request_text: str) -> str:
    # Step 1: Validate the request
    validation = validate_request(request_text)

    # Step 2: Save the request to the database
    record_id = save_request_to_db(user_id, request_text, validation)

    # Step 3: Return a durable result
    return f"Request processed successfully. Record ID: {record_id}"


@step()
def validate_request(text: str) -> str:
    if len(text) < 10:
        return "too short"
    if "urgent" in text.lower():
        return "flagged: urgent"
    return "valid"


@step()
def save_request_to_db(user_id: int, text: str, validation: str) -> int:
    # This is where you'd normally use Django ORM
    # For now, simulate a DB insert
    print(f"Saving request for user {user_id}: {text} ({validation})")
    return 12345  # pretend DB primary key

import os
from datetime import datetime
from docx import Document
from docx2pdf import convert

OUTPUT_DIR = "generated_docs"

@workflow()
def generate_coop_document(coordinator: str, division: str, details: str) -> dict:
    # Step 1: Ensure output directory exists
    ensure_output_directory()

    # Step 2: Create the Word document
    docx_path = create_word_doc(coordinator, division, details)

    # Step 3: Convert to PDF
    pdf_path = convert_to_pdf(docx_path)

    return {
        "word_file": docx_path,
        "pdf_file": pdf_path,
    }


@step()
def ensure_output_directory():
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)


@step()
def create_word_doc(coordinator: str, division: str, details: str) -> str:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"coop_request_{timestamp}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)

    doc = Document()
    doc.add_heading("COOP Request Summary", level=1)

    doc.add_paragraph(f"Coordinator: {coordinator}")
    doc.add_paragraph(f"Division: {division}")
    doc.add_paragraph("Details:")
    doc.add_paragraph(details)

    doc.save(filepath)
    return filepath


@step()
def convert_to_pdf(docx_path: str) -> str:
    pdf_path = docx_path.replace(".docx", ".pdf")
    convert(docx_path, pdf_path)
    return pdf_path

