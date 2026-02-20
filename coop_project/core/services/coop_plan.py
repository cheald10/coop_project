"""
app/services/coop_plan.py

Generates a Word (.docx) COOP plan from all division data,
saves it to MEDIA_ROOT/plans/, and records a GeneratedPlan entry.
PDF conversion is attempted via LibreOffice if available.
"""
import os
import subprocess
import logging
from datetime import date

from django.conf import settings
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.models import (
    Division, EssentialFunction, CriticalApplication, KeyPersonnel,
    VitalRecord, Dependency, AlternateFacility, Communication,
    RecoveryPriority, DivisionMetadata, GeneratedPlan
)

logger = logging.getLogger(__name__)


# ----------------------------------------------------------------
# Public entry point
# ----------------------------------------------------------------

def generate_coop_plan_for_division(division_id: int, generated_by=None) -> dict:
    """
    Generates Word + optional PDF for the given division.

    Returns:
        {
            "success": bool,
            "docx_path": str,
            "pdf_path": str | None,
            "version": int,
            "error": str | None,
        }
    """
    try:
        division = Division.objects.get(pk=division_id)
    except Division.DoesNotExist:
        return {"success": False, "error": f"Division {division_id} not found."}

    # Bump version
    division.plan_version += 1
    division.save(update_fields=["plan_version"])
    version = division.plan_version

    # Ensure output directory exists
    output_dir = os.path.join(settings.MEDIA_ROOT, "plans")
    os.makedirs(output_dir, exist_ok=True)

    safe_name = division.name.replace(" ", "_").replace("/", "-")
    filename_base = f"{safe_name}_COOP_v{version}"
    docx_path = os.path.join(output_dir, f"{filename_base}.docx")
    pdf_path = None

    try:
        doc = _build_document(division)
        doc.save(docx_path)
        logger.info("Saved DOCX: %s", docx_path)

        pdf_path = _convert_to_pdf(docx_path, output_dir)

        # Record in history
        GeneratedPlan.objects.create(
            division=division,
            version=version,
            docx_path=docx_path,
            pdf_path=pdf_path or "",
            created_by=generated_by,
        )

        return {
            "success": True,
            "docx_path": docx_path,
            "pdf_path": pdf_path,
            "version": version,
            "error": None,
        }

    except Exception as exc:
        logger.exception("Plan generation failed for division %s", division_id)
        return {
            "success": False,
            "docx_path": docx_path if os.path.exists(docx_path) else None,
            "pdf_path": None,
            "version": version,
            "error": str(exc),
        }


# ----------------------------------------------------------------
# Document builder
# ----------------------------------------------------------------

def _build_document(division: Division) -> Document:
    doc = Document()

    # ---- Cover page ----
    _heading(doc, f"CONTINUITY OF OPERATIONS PLAN", level=0)
    _heading(doc, division.name, level=1)
    doc.add_paragraph(f"Version {division.plan_version}  |  {date.today():%B %d, %Y}")
    doc.add_page_break()

    # ---- Division Profile ----
    try:
        meta = division.divisionmetadata
        _heading(doc, "Division Profile", level=1)
        _table_two_col(doc, [
            ("Director", meta.director),
            ("Coordinator", meta.coordinator),
            ("Primary Location", meta.primary_location),
            ("Staff Count", str(meta.staff_count)),
            ("Hours of Operation", meta.hours_of_operation),
            ("Plan Status", division.plan_status),
            ("Next Review Date", str(division.next_review_date or "TBD")),
        ])
        if meta.mission_statement:
            _heading(doc, "Mission Statement", level=2)
            doc.add_paragraph(meta.mission_statement)
        if meta.critical_services_summary:
            _heading(doc, "Critical Services Summary", level=2)
            doc.add_paragraph(meta.critical_services_summary)
        doc.add_page_break()
    except DivisionMetadata.DoesNotExist:
        pass

    # ---- Essential Functions ----
    functions = EssentialFunction.objects.filter(division=division)
    if functions.exists():
        _heading(doc, "Essential Functions", level=1)
        headers = ["Function", "Priority", "MTD", "RTO", "Owner"]
        rows = [[f.name, f.priority, f.mtd, f.rto, f.owner] for f in functions]
        _table(doc, headers, rows)
        doc.add_page_break()

    # ---- Critical Applications ----
    apps = CriticalApplication.objects.filter(division=division)
    if apps.exists():
        _heading(doc, "Critical Applications", level=1)
        headers = ["Application", "Hosting", "RTO", "Recovery Tier", "Vendor Contact"]
        rows = [[a.name, a.hosting_environment, a.rto, a.recovery_tier, a.vendor_contact] for a in apps]
        _table(doc, headers, rows)
        doc.add_page_break()

    # ---- Key Personnel ----
    personnel = KeyPersonnel.objects.filter(division=division)
    if personnel.exists():
        _heading(doc, "Key Personnel", level=1)
        headers = ["Name", "Role", "Type", "Work Phone", "Mobile", "Email"]
        rows = [[p.name, p.role, p.primary_or_alternate, p.work_phone, p.mobile_phone, p.email] for p in personnel]
        _table(doc, headers, rows)
        doc.add_page_break()

    # ---- Vital Records ----
    records = VitalRecord.objects.filter(division=division)
    if records.exists():
        _heading(doc, "Vital Records", level=1)
        headers = ["Record", "Type", "Format", "Storage Location", "Backup Location", "Priority"]
        rows = [[r.name, r.record_type, r.format, r.storage_location, r.backup_location, r.priority] for r in records]
        _table(doc, headers, rows)
        doc.add_page_break()

    # ---- Dependencies ----
    deps = Dependency.objects.filter(division=division)
    if deps.exists():
        _heading(doc, "Dependencies", level=1)
        headers = ["Dependency", "Type", "Criticality", "Vendor Contact"]
        rows = [[d.name, d.dependency_type, d.criticality, d.vendor_contact] for d in deps]
        _table(doc, headers, rows)
        doc.add_page_break()

    # ---- Alternate Facilities ----
    facilities = AlternateFacility.objects.filter(division=division)
    if facilities.exists():
        _heading(doc, "Alternate Facilities", level=1)
        headers = ["Facility", "Type", "Address", "Capacity", "IT Available", "Contact"]
        rows = [[f.name, f.facility_type, f.address, str(f.capacity), f.it_availability, f.contact] for f in facilities]
        _table(doc, headers, rows)
        doc.add_page_break()

    # ---- Communications ----
    comms = Communication.objects.filter(division=division)
    if comms.exists():
        _heading(doc, "Communications", level=1)
        headers = ["Type", "Method", "Primary Contact", "Backup Contact"]
        rows = [[c.communication_type, c.method, c.primary_contact, c.backup_contact] for c in comms]
        _table(doc, headers, rows)
        doc.add_page_break()

    # ---- Recovery Priorities ----
    priorities = RecoveryPriority.objects.filter(division=division).order_by("priority_level")
    if priorities.exists():
        _heading(doc, "Recovery Priorities", level=1)
        headers = ["Priority", "Item", "Type", "Rationale"]
        rows = [[str(p.priority_level), p.item_name, p.item_type, p.rationale] for p in priorities]
        _table(doc, headers, rows)

    return doc


# ----------------------------------------------------------------
# Helpers
# ----------------------------------------------------------------

def _heading(doc, text, level=1):
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(18)
    else:
        doc.add_heading(text, level=level)


def _table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    for row_data, row in zip(rows, table.rows[1:]):
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text) if cell_text else ""


def _table_two_col(doc, pairs):
    table = doc.add_table(rows=len(pairs), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(pairs):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[i].cells[1].text = str(value) if value else ""


def _convert_to_pdf(docx_path: str, output_dir: str) -> str | None:
    """
    Attempt PDF conversion via LibreOffice.
    Returns PDF path on success, None if LibreOffice is unavailable.
    """
    try:
        result = subprocess.run(
            [
                "libreoffice", "--headless", "--convert-to", "pdf",
                "--outdir", output_dir, docx_path,
            ],
            capture_output=True,
            timeout=60,
        )
        if result.returncode == 0:
            pdf_name = os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
            pdf_path = os.path.join(output_dir, pdf_name)
            if os.path.exists(pdf_path):
                logger.info("PDF created: %s", pdf_path)
                return pdf_path
    except (FileNotFoundError, subprocess.TimeoutExpired) as exc:
        logger.warning("PDF conversion skipped: %s", exc)
    return None
